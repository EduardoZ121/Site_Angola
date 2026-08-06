#!/usr/bin/env python3
"""Generate PDF + DOCX for Kuteka legal/help documents from Markdown sources."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
)

ROOT = Path(__file__).resolve().parents[1]
SOURCES = [
    ROOT / "docs/legal/TERMOS_UTILIZACAO_v1.md",
    ROOT / "docs/legal/POLITICA_PRIVACIDADE_v1.md",
    ROOT / "docs/legal/POLITICA_COOKIES_v1.md",
    ROOT / "docs/help/MANUAL_UTILIZADOR_v1.md",
]
OUT_DOCS = ROOT / "docs" / "legal" / "exports"
OUT_PUBLIC = ROOT / "apps/web/public/docs"


def parse_blocks(md: str) -> list[tuple[str, object]]:
    lines = md.replace("\r\n", "\n").split("\n")
    blocks: list[tuple[str, object]] = []
    i = 0
    para: list[str] = []

    def flush_para() -> None:
        nonlocal para
        if para:
            blocks.append(("p", " ".join(para).strip()))
            para = []

    while i < len(lines):
        line = lines[i]
        if line.strip() == "---":
            flush_para()
            blocks.append(("hr", None))
            i += 1
            continue
        if line.startswith("#"):
            flush_para()
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            blocks.append((f"h{min(level, 3)}", text))
            i += 1
            continue
        if line.startswith("> "):
            flush_para()
            quote = [line[2:]]
            i += 1
            while i < len(lines) and lines[i].startswith("> "):
                quote.append(lines[i][2:])
                i += 1
            blocks.append(("quote", " ".join(quote).strip()))
            continue
        if line.startswith("|") and "|" in line[1:]:
            flush_para()
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                raw = lines[i].strip()
                if re.match(r"^\|[\s|:-]+\|$", raw):
                    i += 1
                    continue
                cells = [c.strip() for c in raw.strip("|").split("|")]
                rows.append(cells)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue
        if re.match(r"^[-*]\s+", line):
            flush_para()
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i]):
                items.append(re.sub(r"^[-*]\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("ul", items))
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_para()
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]).strip())
                i += 1
            blocks.append(("ol", items))
            continue
        if not line.strip():
            flush_para()
            i += 1
            continue
        para.append(line.strip())
        i += 1
    flush_para()
    return blocks


def md_inline(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`(.+?)`", r"<font face='Courier'>\1</font>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def plain_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def build_pdf(blocks: list[tuple[str, object]], dest: Path, title: str) -> None:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="KTitle",
            parent=styles["Title"],
            fontSize=16,
            leading=20,
            spaceAfter=12,
            textColor=colors.HexColor("#08263f"),
            alignment=TA_CENTER,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KH1",
            parent=styles["Heading1"],
            fontSize=13,
            leading=17,
            spaceBefore=14,
            spaceAfter=6,
            textColor=colors.HexColor("#08263f"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="KH2",
            parent=styles["Heading2"],
            fontSize=11,
            leading=14,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor("#0f3a5c"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="KH3",
            parent=styles["Heading3"],
            fontSize=10,
            leading=13,
            spaceBefore=8,
            spaceAfter=3,
            textColor=colors.HexColor("#1a4a6e"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="KBody",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KQuote",
            parent=styles["BodyText"],
            fontSize=8.5,
            leading=11,
            leftIndent=10,
            textColor=colors.HexColor("#334155"),
            spaceAfter=8,
            spaceBefore=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KMeta",
            parent=styles["BodyText"],
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#475569"),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="KCell",
            parent=styles["BodyText"],
            fontSize=7.5,
            leading=9.5,
        )
    )

    doc = SimpleDocTemplate(
        str(dest),
        pagesize=A4,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title=title,
        author="Kuteka",
    )
    story = []
    for kind, payload in blocks:
        if kind == "h1":
            story.append(Paragraph(md_inline(str(payload)), styles["KTitle"]))
        elif kind == "h2":
            story.append(Paragraph(md_inline(str(payload)), styles["KH1"]))
        elif kind == "h3":
            story.append(Paragraph(md_inline(str(payload)), styles["KH2"]))
        elif kind == "p":
            style = styles["KMeta"] if payload == blocks[1][1] and kind == "p" else styles["KBody"]
            # meta lines often start with **Versão
            if str(payload).startswith("**Versão") or str(payload).startswith("**Manual"):
                story.append(Paragraph(md_inline(str(payload)), styles["KMeta"]))
            else:
                story.append(Paragraph(md_inline(str(payload)), styles["KBody"]))
        elif kind == "quote":
            story.append(Paragraph(md_inline(str(payload)), styles["KQuote"]))
        elif kind == "hr":
            story.append(Spacer(1, 6))
        elif kind == "ul":
            for item in payload:  # type: ignore[union-attr]
                story.append(Paragraph(f"• {md_inline(item)}", styles["KBody"]))
        elif kind == "ol":
            for idx, item in enumerate(payload, 1):  # type: ignore[arg-type]
                story.append(Paragraph(f"{idx}. {md_inline(item)}", styles["KBody"]))
        elif kind == "table":
            rows = payload  # type: ignore[assignment]
            data = [
                [Paragraph(md_inline(c), styles["KCell"]) for c in row] for row in rows  # type: ignore[union-attr]
            ]
            col_count = max(len(r) for r in rows)  # type: ignore[arg-type]
            width = 17 * cm / col_count
            table = Table(data, colWidths=[width] * col_count, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(KeepTogether([table, Spacer(1, 8)]))
    doc.build(story)


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def build_docx(blocks: list[tuple[str, object]], dest: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, payload in blocks:
        if kind == "h1":
            p = doc.add_heading(plain_inline(str(payload)), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "h2":
            doc.add_heading(plain_inline(str(payload)), level=1)
        elif kind == "h3":
            doc.add_heading(plain_inline(str(payload)), level=2)
        elif kind == "p":
            p = doc.add_paragraph()
            add_runs(p, str(payload))
            if str(payload).startswith("**Versão") or str(payload).startswith("**Manual"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        elif kind == "quote":
            p = doc.add_paragraph()
            add_runs(p, str(payload))
            p.paragraph_format.left_indent = Pt(12)
            for run in p.runs:
                run.italic = True
        elif kind == "ul":
            for item in payload:  # type: ignore[union-attr]
                p = doc.add_paragraph(style="List Bullet")
                p.clear()
                add_runs(p, item)
        elif kind == "ol":
            for item in payload:  # type: ignore[union-attr]
                p = doc.add_paragraph(style="List Number")
                p.clear()
                add_runs(p, item)
        elif kind == "table":
            rows = payload  # type: ignore[assignment]
            cols = max(len(r) for r in rows)  # type: ignore[arg-type]
            table = doc.add_table(rows=len(rows), cols=cols)  # type: ignore[arg-type]
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):  # type: ignore[arg-type]
                for c_idx in range(cols):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    table.rows[r_idx].cells[c_idx].text = plain_inline(cell_text)
            doc.add_paragraph("")
        elif kind == "hr":
            doc.add_paragraph("—")

    doc.save(str(dest))


def main() -> None:
    OUT_DOCS.mkdir(parents=True, exist_ok=True)
    OUT_PUBLIC.mkdir(parents=True, exist_ok=True)

    for src in SOURCES:
        if not src.exists():
            raise SystemExit(f"Missing source: {src}")
        md = src.read_text(encoding="utf-8")
        blocks = parse_blocks(md)
        stem = src.stem
        title = plain_inline(next((p for k, p in blocks if k == "h1"), stem))

        pdf_path = OUT_DOCS / f"{stem}.pdf"
        docx_path = OUT_DOCS / f"{stem}.docx"
        build_pdf(blocks, pdf_path, title)
        build_docx(blocks, docx_path)

        # Public downloads (md + pdf + docx)
        shutil.copy2(src, OUT_PUBLIC / f"{stem}.md")
        shutil.copy2(pdf_path, OUT_PUBLIC / f"{stem}.pdf")
        shutil.copy2(docx_path, OUT_PUBLIC / f"{stem}.docx")
        # Also mirror under docs/help exports for manual
        if "MANUAL" in stem:
            help_export = ROOT / "docs/help/exports"
            help_export.mkdir(parents=True, exist_ok=True)
            shutil.copy2(pdf_path, help_export / f"{stem}.pdf")
            shutil.copy2(docx_path, help_export / f"{stem}.docx")

        print(f"OK {stem}: {len(blocks)} blocks → pdf/docx/md")


if __name__ == "__main__":
    main()
